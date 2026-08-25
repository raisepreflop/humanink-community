#!/usr/bin/env python3
"""
AWOS md2docx v2 — Conversor Markdown → Word con formato predefinido + track changes + versionado

Formato predefinido:
  Times New Roman 12, 1.5 espaciado, justificado, sangría primera línea
  Título 1: 14pt Bold, salto de página   |   Título 2: 13pt Bold
  Separadores: *** centrado
  Páginas: A4, márgenes 1 pulgada

Track changes:
  Las inserciones se marcan con <w:ins> (verde en Word)
  Las eliminaciones se marcan con <w:del> (tachado rojo en Word)

Versionado:
  get_next_version(base_path) → detecta vN y devuelve N+1
"""

import sys
import re
import os
import subprocess
from pathlib import Path
from datetime import datetime

TRACK_AUTHOR = "AWOS Escritor Fantasma"


def install_deps():
    try:
        import docx  # noqa
    except ImportError:
        subprocess.run(
            [sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'],
            capture_output=True
        )


install_deps()

from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────
# VERSIONADO
# ──────────────────────────────────────────────

def get_next_version(output_path: str) -> tuple[str, int]:
    """
    A partir de 'capitulos/cap-01.docx' o 'capitulos/cap-01-v3.docx'
    detecta la última versión existente y devuelve (ruta_nueva, numero_version).
    Ejemplo: cap-01-v3.docx existe → devuelve ('cap-01-v4.docx', 4)
    """
    p = Path(output_path)
    stem = p.stem
    parent = p.parent
    ext = p.suffix  # '.docx'

    # Quitar sufijo -vN si ya existe
    base_stem = re.sub(r'-v\d+$', '', stem)

    # Buscar versiones existentes
    existing = sorted(
        parent.glob(f'{base_stem}-v*.docx'),
        key=lambda f: int(re.search(r'-v(\d+)', f.stem).group(1)) if re.search(r'-v(\d+)', f.stem) else 0
    )
    if not existing:
        # Primera versión
        new_path = str(parent / f'{base_stem}-v1{ext}')
        return new_path, 1

    last_v = int(re.search(r'-v(\d+)', existing[-1].stem).group(1))
    new_v = last_v + 1
    new_path = str(parent / f'{base_stem}-v{new_v}{ext}')
    return new_path, new_v


# ──────────────────────────────────────────────
# ESTILOS
# ──────────────────────────────────────────────

def setup_styles(doc):
    """Configura los estilos Word con el formato predefinido."""

    n = doc.styles['Normal']
    n.font.name = 'Times New Roman'
    n.font.size = Pt(12)
    pf = n.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(18)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = None
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.page_break_before = False
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = None
    h3.paragraph_format.page_break_before = False
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = Pt(0)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


# ──────────────────────────────────────────────
# INLINE FORMATTING
# ──────────────────────────────────────────────

def add_inline(paragraph, text):
    """Añade texto con soporte de **negrita** e *cursiva* inline."""
    pattern = re.compile(r'(\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
        elif part:
            paragraph.add_run(part)


def is_bold_only_line(line):
    stripped = line.strip()
    return (stripped.startswith('**') and stripped.endswith('**')
            and stripped.count('**') == 2 and len(stripped) > 4)


# ──────────────────────────────────────────────
# TRACK CHANGES — XML helpers
# ──────────────────────────────────────────────

_rev_counter = [0]


def _next_rev_id():
    _rev_counter[0] += 1
    return _rev_counter[0]


def _tc_date():
    return datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')


def _make_ins_para(doc, text, author=TRACK_AUTHOR, is_bold=False, is_center=False):
    """
    Añade un párrafo al documento con su contenido envuelto en <w:ins>.
    Devuelve el elemento <w:p> resultante.
    """
    p_elem = OxmlElement('w:p')

    # Propiedades de párrafo
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)

    if is_center:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        pPr.append(ind)

    # rPr change mark
    pPrChange = OxmlElement('w:rPr')
    rPrChange_ins = OxmlElement('w:ins')
    rPrChange_ins.set(qn('w:id'), str(_next_rev_id()))
    rPrChange_ins.set(qn('w:author'), author)
    rPrChange_ins.set(qn('w:date'), _tc_date())
    pPrChange.append(rPrChange_ins)
    pPr.append(pPrChange)
    p_elem.append(pPr)

    # Insertion mark wrapping the run
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(_next_rev_id()))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), _tc_date())

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if is_bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
    rPr.append(sz)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    ins.append(r)
    p_elem.append(ins)

    doc.element.body.append(p_elem)
    return p_elem


def _make_del_para(doc, text, author=TRACK_AUTHOR):
    """
    Añade un párrafo con su contenido envuelto en <w:del> (tachado/eliminado).
    """
    p_elem = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p_elem.append(pPr)

    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), str(_next_rev_id()))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), _tc_date())

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    r.append(rPr)

    dt = OxmlElement('w:delText')
    dt.text = text
    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(dt)
    del_elem.append(r)
    p_elem.append(del_elem)

    doc.element.body.append(p_elem)
    return p_elem


# ──────────────────────────────────────────────
# CONVERSOR PRINCIPAL: MD → DOCX
# ──────────────────────────────────────────────

def md_to_docx(md_content, output_path, doc_title='', tracked=False, author=TRACK_AUTHOR):
    """
    Convierte markdown a .docx con formato predefinido.
    tracked=True → todo el contenido se marca como <w:ins> (inserción nueva).
    """
    doc = Document()

    section = doc.sections[0]
    # A4 con unidades explícitas de python-docx: un entero crudo se interpreta como
    # EMU (914400/pulgada), no twips — 11906 EMU = 0.013" (bug real reportado por un usuario:
    # "un carácter por página").
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    setup_styles(doc)

    if doc_title:
        p = doc.add_heading(doc_title, level=1)
        p.paragraph_format.page_break_before = False

    in_code_block = False
    lines = md_content.split('\n')

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        if not stripped:
            continue

        if tracked:
            # En modo tracked, añadir todo como inserción
            is_sep = stripped in ('---', '***', '* * *')
            is_h = stripped.startswith('#')
            if is_h:
                level = len(stripped) - len(stripped.lstrip('#'))
                text = stripped.lstrip('#').strip()
                _make_ins_para(doc, text, author=author, is_bold=True)
            elif is_sep:
                _make_ins_para(doc, '***', author=author, is_center=True)
            elif is_bold_only_line(stripped):
                _make_ins_para(doc, stripped[2:-2], author=author, is_bold=True)
            else:
                # Strip simple markdown
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
                clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
                _make_ins_para(doc, clean, author=author)
        else:
            # Modo normal — full formatting
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:].strip(), level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith('# '):
                p = doc.add_heading(stripped[2:].strip(), level=1)
                p.paragraph_format.page_break_before = False
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:].strip(), level=3)
            elif stripped in ('---', '***', '* * *', '— — —'):
                p = doc.add_paragraph()
                p.add_run('***')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            elif stripped.startswith('> ') or stripped == '>':
                text = stripped[2:].strip() if stripped.startswith('> ') else ''
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Pt(36)
                p.paragraph_format.first_line_indent = Pt(0)
                p.add_run(text).italic = True
            elif is_bold_only_line(stripped):
                doc.add_heading(stripped[2:-2], level=3)
            elif re.match(r'^[-*•]\s+', stripped):
                text = re.sub(r'^[-*•]\s+', '', stripped)
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Pt(36)
                p.add_run('• ')
                add_inline(p, text)
            elif re.match(r'^\d+[.)]\s+', stripped):
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Pt(36)
                add_inline(p, stripped)
            else:
                p = doc.add_paragraph(style='Normal')
                add_inline(p, stripped)

    doc.save(output_path)
    return output_path


# ──────────────────────────────────────────────
# MERGE CON TRACK CHANGES (para reescrituras y revisiones)
# ──────────────────────────────────────────────

def create_tracked_revision(base_docx_path, new_md_content, output_path,
                             mode='rewrite', section_marker=None,
                             author=TRACK_AUTHOR, version=None):
    """
    Crea un nuevo .docx a partir del original con track changes:

    mode='rewrite'   → todo el original marcado como eliminado + nuevo como insertado
    mode='insert'    → contenido original sin tocar + nuevo fragmento insertado
                       al final (o después de section_marker si se especifica)
    mode='section'   → identifica la sección por section_marker, la marca como
                       eliminada, inserta la nueva versión en su lugar

    version: número de versión para la nota de revisión.
    """
    base = Document(base_docx_path)
    out = Document()

    # Copiar configuración de página del original
    s = out.sections[0]
    # Mismas unidades explícitas que arriba (el bug de twips-como-EMU estaba duplicado aquí).
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    setup_styles(out)

    # Cabecera de revisión
    date_str = datetime.now().strftime('%d/%m/%Y %H:%M')
    v_label = f'v{version}' if version else ''
    note_text = f'[REVISIÓN {v_label} — {date_str} — {author}]'
    p_note = out.add_paragraph(style='Normal')
    r_note = p_note.add_run(note_text)
    r_note.bold = True
    r_note.italic = True
    p_note.paragraph_format.first_line_indent = Pt(0)

    sep_p = out.add_paragraph()
    sep_p.add_run('***')
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_p.paragraph_format.first_line_indent = Pt(0)

    base_paragraphs = [p.text for p in base.paragraphs if p.text.strip()]

    if mode == 'rewrite':
        # Todo el original → eliminado
        for para_text in base_paragraphs:
            _make_del_para(out, para_text, author=author)

        # Separador
        sep_p2 = out.add_paragraph()
        sep_p2.add_run('***')
        sep_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_p2.paragraph_format.first_line_indent = Pt(0)

        # Nuevo contenido → insertado
        _inject_tracked_md(out, new_md_content, author=author)

    elif mode == 'insert':
        # Original sin tocar
        for para_text in base_paragraphs:
            p = out.add_paragraph(para_text, style='Normal')

        # Separador antes del fragmento nuevo
        sep_p2 = out.add_paragraph()
        sep_p2.add_run('***')
        sep_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_p2.paragraph_format.first_line_indent = Pt(0)

        # Nuevo fragmento → insertado (track changes)
        _inject_tracked_md(out, new_md_content, author=author)

    elif mode == 'section':
        # Identificar sección a reemplazar
        inside_section = False
        section_found = False
        buffered_del = []

        for para_text in base_paragraphs:
            if section_marker and section_marker.lower() in para_text.lower():
                inside_section = True
                section_found = True
                buffered_del.append(para_text)
                continue

            if inside_section:
                # Terminamos la sección si encontramos el siguiente marcador de sección
                if (para_text.strip().startswith('###') or
                        para_text.strip().startswith('##') or
                        para_text.strip() in ('---', '***')):
                    inside_section = False
                    # Vaciar el buffer como eliminados
                    for del_text in buffered_del:
                        _make_del_para(out, del_text, author=author)
                    buffered_del = []
                    # Insertar nuevo contenido en su lugar
                    _inject_tracked_md(out, new_md_content, author=author)
                    # Añadir el párrafo actual (siguiente sección) sin cambios
                    out.add_paragraph(para_text, style='Normal')
                else:
                    buffered_del.append(para_text)
            else:
                out.add_paragraph(para_text, style='Normal')

        # Si la sección llegó hasta el final
        if inside_section and buffered_del:
            for del_text in buffered_del:
                _make_del_para(out, del_text, author=author)
            _inject_tracked_md(out, new_md_content, author=author)

        if not section_found:
            # Si no se encontró el marcador, insertar al final
            sep_p3 = out.add_paragraph()
            sep_p3.add_run('*** SECCIÓN NO ENCONTRADA — FRAGMENTO AÑADIDO AL FINAL ***')
            sep_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep_p3.paragraph_format.first_line_indent = Pt(0)
            _inject_tracked_md(out, new_md_content, author=author)

    out.save(output_path)
    return output_path


def _inject_tracked_md(doc, md_content, author=TRACK_AUTHOR):
    """Inyecta contenido markdown como párrafos con track changes (inserción)."""
    lines = md_content.split('\n')
    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('```'):
            in_code = not in_code
            continue
        if in_code or not stripped:
            continue
        is_sep = stripped in ('---', '***', '* * *')
        is_heading = stripped.startswith('#')
        if is_heading:
            text = stripped.lstrip('#').strip()
            _make_ins_para(doc, text, author=author, is_bold=True)
        elif is_sep:
            _make_ins_para(doc, '***', author=author, is_center=True)
        else:
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            _make_ins_para(doc, clean, author=author)


# ──────────────────────────────────────────────
# LEER DOCX
# ──────────────────────────────────────────────

def read_docx_text(docx_path):
    """Extrae el texto plano de un .docx."""
    doc = Document(docx_path)
    return '\n\n'.join(p.text for p in doc.paragraphs if p.text.strip())


# ──────────────────────────────────────────────
# INSTALADOR
# ──────────────────────────────────────────────

def install_self():
    import shutil
    awos_dir = os.path.expanduser('~/.awos')
    os.makedirs(awos_dir, exist_ok=True)
    dest = os.path.join(awos_dir, 'md2docx.py')
    shutil.copy2(__file__, dest)
    print(f'✓ Instalado en: {dest}')


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='AWOS md2docx v2')
    parser.add_argument('input', nargs='?', help='Fichero .md de entrada')
    parser.add_argument('output', nargs='?', help='Fichero .docx de salida')
    parser.add_argument('title', nargs='?', default='', help='Título del documento')
    parser.add_argument('--install', action='store_true', help='Instalar en ~/.awos/')
    parser.add_argument('--tracked', action='store_true', help='Marcar todo como inserción')
    parser.add_argument('--version', action='store_true', help='Incrementar versión automáticamente')
    parser.add_argument('--base', help='Docx base para track changes')
    parser.add_argument('--mode', choices=['rewrite', 'insert', 'section'], default='rewrite')
    parser.add_argument('--section-marker', help='Texto que identifica la sección a reemplazar')
    args = parser.parse_args()

    if args.install:
        install_self()
        sys.exit(0)

    if not args.input or not args.output:
        print('Uso: md2docx.py input.md output.docx ["Título"]')
        print('     md2docx.py input.md output.docx --version  (auto-versión)')
        print('     md2docx.py input.md output.docx --base orig.docx --mode rewrite')
        print('     md2docx.py --install')
        sys.exit(1)

    with open(args.input, 'r', encoding='utf-8') as f:
        content = f.read()

    output = args.output
    version_num = None

    if args.version:
        output, version_num = get_next_version(args.output)
        print(f'Versión detectada: v{version_num} → {output}')

    if args.base and os.path.exists(args.base):
        result = create_tracked_revision(
            base_docx_path=args.base,
            new_md_content=content,
            output_path=output,
            mode=args.mode,
            section_marker=args.section_marker,
            version=version_num
        )
    else:
        result = md_to_docx(content, output, args.title, tracked=args.tracked)

    print(f'✓ Word guardado: {result}')
    if version_num:
        print(f'  Versión: v{version_num}')
