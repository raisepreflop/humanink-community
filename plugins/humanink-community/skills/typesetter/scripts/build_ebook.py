#!/usr/bin/env python3
"""
build_ebook.py — Pipeline reproducible DOCX (track changes) -> EPUB validado para Amazon KDP.

Genérico: sirve para cualquier libro con la convención de estilos «de casa»
(Heading 1 = partes, Heading 2 = secciones/capítulos, separador de escena «* * *»).

Uso típico (desde la carpeta del libro):
    python3 build_ebook.py --src TAW-1-b25.docx

Opciones:
    --src        manuscrito .docx (relativo a --book-dir o absoluto)   [requerido]
    --book-dir   carpeta del libro (def: directorio actual)
    --cover      portada JPG (def: covers/kindle-cover-kdp.jpg en book-dir)
    --css        hoja de estilo (def: assets/ebook.css del libro, o la del skill)
    --metadata   metadata.yaml (def: assets/metadata.yaml del libro, o la del skill)
    --out        epub de salida (def: book-dir/output/<slug>.epub)
    --no-validate  omitir EPUBCheck

Flujo:
  1. Extrae el texto ACEPTADO (incluye <w:ins>, ignora <w:del>/<w:delText>),
     preservando cursiva/negrita de los runs.
  2. DESCARTA la portada/créditos/«Contents» manual del .docx (todo lo anterior
     al primer encabezado): pandoc genera la portada y la página de título desde metadata.
  3. Mapea Heading 1 + Heading 2 -> '# ' (h1) => índice de navegación plano.
     '* * *' -> ::: scene (div centrado). Tablas -> '**término** — descripción'.
  4. pandoc -> EPUB con CSS y metadata versionados.
  5. Valida con EPUBCheck; aborta si hay errores/fatales.

Requisitos: python-docx, PIL, pandoc, epubcheck, java.
  macOS:  brew install pandoc epubcheck   (java viene con epubcheck / Temurin)
          pip3 install python-docx pillow
"""
import argparse
import re
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from PIL import Image
except ImportError as e:
    sys.exit(f"Falta dependencia Python ({e}). Instala: pip3 install python-docx pillow")

SKILL_DIR = Path(__file__).resolve().parent          # .../typesetter/scripts
SKILL_ASSETS = SKILL_DIR / "assets"

W_T, W_TAB, W_BR = qn("w:t"), qn("w:tab"), qn("w:br")
W_R, W_DEL, W_RPR = qn("w:r"), qn("w:del"), qn("w:rPr")
W_I, W_B = qn("w:i"), qn("w:b")
W_P, W_TBL = qn("w:p"), qn("w:tbl")

SCENE_RE = re.compile(r"^[\*∗•·\s]{1,8}$")


def md_escape(text: str) -> str:
    return re.sub(r"([\\`*_<>\[\]])", r"\\\1", text)


def run_is_deleted(run) -> bool:
    return any(anc.tag == W_DEL for anc in run.iterancestors())


def run_text(run) -> str:
    out = []
    for node in run.iter():
        if node.tag == W_T:
            out.append(node.text or "")
        elif node.tag in (W_TAB, W_BR):
            out.append(" ")
    return "".join(out)


def run_style(run):
    rpr = run.find(W_RPR)
    if rpr is None:
        return (False, False)
    def on(tag):
        el = rpr.find(tag)
        return el is not None and el.get(qn("w:val")) not in ("0", "false", "none")
    return (on(W_I), on(W_B))


def paragraph_markdown(p) -> str:
    groups = []
    for run in p._p.iter(W_R):
        if run_is_deleted(run):
            continue
        txt = run_text(run)
        if not txt:
            continue
        ital, bold = run_style(run)
        if groups and groups[-1][0] == ital and groups[-1][1] == bold:
            groups[-1] = (ital, bold, groups[-1][2] + txt)
        else:
            groups.append((ital, bold, txt))
    pieces = []
    for ital, bold, txt in groups:
        esc = md_escape(txt)
        lead = esc[: len(esc) - len(esc.lstrip())]
        trail = esc[len(esc.rstrip()):]
        core = esc.strip()
        if core:
            if ital and bold:
                core = f"***{core}***"
            elif bold:
                core = f"**{core}**"
            elif ital:
                core = f"*{core}*"
        pieces.append(f"{lead}{core}{trail}")
    return "".join(pieces).strip()


def accepted_plain(p) -> str:
    return "".join(run_text(r) for r in p._p.iter(W_R) if not run_is_deleted(r)).strip()


def table_markdown(tbl) -> list:
    lines = []
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        dedup = []
        for c in cells:
            if not dedup or dedup[-1] != c:
                dedup.append(c)
        if not any(dedup):
            continue
        term = md_escape(dedup[0])
        desc = md_escape(" ".join(dedup[1:]).strip()) if len(dedup) > 1 else ""
        lines.append(f"**{term}** — {desc}" if desc else f"**{term}**")
        lines.append("")
    return lines


def build_markdown(src: Path) -> str:
    doc = Document(str(src))
    body = doc.element.body
    tables = {id(t._tbl): t for t in doc.tables}
    paras = {id(p._p): p for p in doc.paragraphs}
    lines = []
    started = False  # ignora portada/créditos/«Contents» del .docx hasta el 1er encabezado
    for child in body.iterchildren():
        if child.tag == W_P:
            p = paras.get(id(child))
            if p is None:
                continue
            style = (p.style.name or "").strip()
            plain = accepted_plain(p)
            if not plain:
                continue
            if not started:
                if style in ("Heading 1", "Heading 2"):
                    started = True
                else:
                    continue
            if style in ("Heading 1", "Heading 2"):
                lines += [f"# {md_escape(plain)}", ""]
            elif SCENE_RE.match(plain) and "*" in plain:
                lines += ["::: scene", "\\* \\* \\*", ":::", ""]
            else:
                lines += [paragraph_markdown(p), ""]
        elif child.tag == W_TBL and started:
            tbl = tables.get(id(child))
            if tbl is not None:
                lines += table_markdown(tbl)
    md = re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip() + "\n"
    return md


def resolve(book_dir: Path, given, rel, fallback=None) -> Path:
    if given:
        p = Path(given)
        return p if p.is_absolute() else (book_dir / p)
    local = book_dir / rel
    if local.exists():
        return local
    return fallback if fallback else local


def run_epubcheck(epub: Path) -> bool:
    print(f"\n→ EPUBCheck {epub.name} …")
    try:
        proc = subprocess.run(["epubcheck", str(epub)], capture_output=True, text=True)
    except FileNotFoundError:
        print("  ⚠ epubcheck no instalado (brew install epubcheck). Validación OMITIDA.")
        return True
    print((proc.stdout + proc.stderr).strip())
    ok = proc.returncode == 0
    print("✓ EPUBCheck: 0 errores" if ok else "✗ EPUBCheck reportó problemas")
    return ok


def main():
    ap = argparse.ArgumentParser(description="DOCX -> EPUB validado para KDP.")
    ap.add_argument("--src", required=True)
    ap.add_argument("--book-dir", default=".")
    ap.add_argument("--cover")
    ap.add_argument("--css")
    ap.add_argument("--metadata")
    ap.add_argument("--out")
    ap.add_argument("--no-validate", action="store_true")
    args = ap.parse_args()

    book = Path(args.book_dir).resolve()
    src = resolve(book, args.src, args.src)
    if not src.exists():
        sys.exit(f"No existe el fuente: {src}")

    cover = resolve(book, args.cover, "covers/kindle-cover-kdp.jpg")
    css = resolve(book, args.css, "assets/ebook.css", SKILL_ASSETS / "ebook.css")
    meta = resolve(book, args.metadata, "assets/metadata.yaml", SKILL_ASSETS / "metadata.yaml")
    for label, p in [("portada", cover), ("css", css), ("metadata", meta)]:
        if not p.exists():
            sys.exit(f"Falta {label}: {p}")

    out_dir = book / "output"
    out_dir.mkdir(exist_ok=True)
    slug = re.sub(r"[^a-z0-9]+", "-", src.stem.lower()).strip("-")
    epub = Path(args.out).resolve() if args.out else out_dir / f"{slug}.epub"
    md_out = out_dir / "manuscrito-completo.md"

    # Portada
    im = Image.open(cover)
    print(f"Portada: {cover.name} {im.size} {im.mode}")
    if im.mode != "RGB":
        print("  ⚠ la portada no es RGB (KDP exige RGB)")
    if min(im.size) < 1600:
        print(f"  ⚠ lado corto {min(im.size)}px < 1600 (KDP recomienda ≥1600)")

    print(f"→ Extrayendo texto aceptado de {src.name} …")
    md = build_markdown(src)
    md_out.write_text(md, encoding="utf-8")
    prose = [ln for ln in md.splitlines()
             if ln.strip() and not ln.startswith(("#", ":::")) and ln.strip() != "\\* \\* \\*"]
    words = len(re.sub(r"[#*\\]", "", " ".join(prose)).split())
    sections = md.count("\n# ") + (1 if md.startswith("# ") else 0)
    print(f"  Markdown: {md_out}  (~{words} palabras, {sections} secciones h1)")

    print("→ pandoc → EPUB …")
    # Compatibilidad de versiones de pandoc:
    #   pandoc ≥3.0 usa --split-level ; pandoc ≤2.x usa --epub-chapter-level
    _help = subprocess.run(["pandoc", "--help"], capture_output=True, text=True).stdout
    split_flag = "--split-level=1" if "--split-level" in _help else "--epub-chapter-level=1"
    cmd = ["pandoc", str(md_out), "-o", str(epub),
           "--css", str(css), "--epub-cover-image", str(cover),
           "--metadata-file", str(meta),
           "--toc", "--toc-depth=1", split_flag]
    if subprocess.run(cmd).returncode != 0:
        sys.exit("pandoc falló")
    print(f"  EPUB: {epub}")

    if not args.no_validate and not run_epubcheck(epub):
        sys.exit(1)
    print(f"\n✅ Listo para KDP: {epub}")


if __name__ == "__main__":
    main()
